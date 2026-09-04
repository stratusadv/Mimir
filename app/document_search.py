# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "openai",
#     "python-dotenv",
#     "python-docx",
# ]
# ///

from __future__ import annotations

import logging
import os
import sys

from pathlib import Path
from time import perf_counter
from types import TracebackType

from docx import Document

from console import Console
from constants import LOG_PATH, SUPPORTED_DOCUMENT_EXTENSIONS
from data import TranscriptionSettings
from document_searcher import DocumentSearcher
from errors import SearchError
from transcript_polisher import TranscriptPolisher


log = logging.getLogger('mimir')


class DocumentSearchManager:
    def __init__(self, settings: TranscriptionSettings, console: Console | None = None) -> None:
        self.settings = settings
        self.console = console if console else Console()
        self.fail_count = 0
        self.output_files: list[Path] = []
        self.success_count = 0

    @staticmethod
    def collect_document_files() -> list[Path]:
        if len(sys.argv) > 1:
            list_file = Path(sys.argv[1])
            lines = list_file.read_text(encoding='utf-8').splitlines()

            return [Path(line.strip()) for line in lines if line.strip()]

        document_files = []

        for document_extension in SUPPORTED_DOCUMENT_EXTENSIONS:
            document_files.extend(list(Path.cwd().glob(f'*.{document_extension}')))

        return document_files

    def process(self, searcher: DocumentSearcher, document_file: Path) -> None:
        start_time = perf_counter()
        self.console.search_start(document_file, self.settings.search_query)
        log.info('search start  file=%s', document_file.name)

        try:
            document_text = document_text_read(document_file)
            findings = searcher.search(document_text, self.settings.search_query)

        except Exception as error:
            log.exception('search failed  file=%s', document_file.name)
            self.console.failure(document_file, error)
            self.fail_count += 1

        else:
            elapsed_seconds = perf_counter() - start_time
            self.process_result(searcher, document_file, findings, elapsed_seconds)

    def process_result(
            self,
            searcher: DocumentSearcher,
            document_file: Path,
            findings: str,
            elapsed_seconds: float,
    ) -> None:
        self.console.search_ok(document_file, elapsed_seconds)
        self.success_count += 1

        output_text = searcher.format_search(
            document_file,
            self.settings.search_query,
            findings,
        )

        self.console.search_findings(output_text)

        if not self.console.search_save_ask():
            self.console.search_skipped()
            log.info('search ok  file=%s  saved=no  seconds=%.0f', document_file.name, elapsed_seconds)

            return

        output_file = searcher.search_file_write(
            document_file,
            self.settings.search_query,
            findings,
        )

        self.output_files.append(output_file)
        self.console.search_saved(output_file)

        log.info(
            'search ok  file=%s  output=%s  seconds=%.0f',
            document_file.name,
            output_file.name,
            elapsed_seconds,
        )

    def report_result_file(self) -> None:
        if not (self.settings.result_file and self.output_files):
            return

        lines = '\n'.join(str(output_file) for output_file in self.output_files)
        self.settings.result_file.write_text(lines, encoding='utf-8')

    def run(self) -> int:
        if not self.settings.is_configured:
            detail = self.settings.configuration_detail
            self.console.error('No API settings were found.', detail)
            log.error('no API settings were found. %s', detail)

            return 1

        if not self.settings.search_query:
            self.console.error('No search request was given.')
            log.error('no search request was given')

            return 1

        document_files = self.collect_document_files()
        log.info('queued %s file(s)', len(document_files))

        if not document_files:
            self.console.warning('No documents to search.')
            log.warning('no documents to search')

            return 0

        searcher = DocumentSearcher(TranscriptPolisher(self.settings))

        for document_file in document_files:
            self.process(searcher, document_file)

        self.console.summary(self.success_count, self.fail_count, 0, 0, success_label='searched')
        self.report_result_file()

        return 0


def configure_logging() -> None:
    if log.handlers:
        return

    if LOG_PATH.exists() and LOG_PATH.stat().st_size:
        with LOG_PATH.open('a', encoding='utf-8') as log_file:
            log_file.write('\n')

    formatter = logging.Formatter(
        fmt='%(asctime)s  %(levelname)-8s  %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S',
    )

    try:
        handler = logging.FileHandler(LOG_PATH, encoding='utf-8')

    except OSError:
        print(f'   could not write log file: {LOG_PATH}', file=sys.stderr)

        return

    handler.setFormatter(formatter)
    log.addHandler(handler)
    log.setLevel(logging.INFO)
    log.propagate = False
    logging.captureWarnings(True)
    logging.getLogger().addHandler(handler)
    logging.getLogger().setLevel(logging.WARNING)
    sys.excepthook = unhandled_exception


def document_docx_read(document_file: Path) -> str:
    document = Document(str(document_file))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            lines.append(paragraph_text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = ' | '.join(cell for cell in cells if cell)

            if line:
                lines.append(line)

    return '\n'.join(lines)


def document_text_read(document_file: Path) -> str:
    suffix = document_file.suffix[1:].lower()

    if suffix not in SUPPORTED_DOCUMENT_EXTENSIONS:
        message = (
            f'{document_file.name} has unsupported extension, '
            f'choices are: {SUPPORTED_DOCUMENT_EXTENSIONS}'
        )

        raise SearchError(message)

    if suffix == 'docx':
        return document_docx_read(document_file)

    return document_file.read_text(encoding='utf-8-sig')


def main() -> int:
    configure_logging()
    Console.enable_ansi_colors()
    settings = TranscriptionSettings.from_environment()
    manager = DocumentSearchManager(settings)

    log.info(
        'search session start  python=%s  text_model=%s  host=%s  key=%s',
        sys.version.split()[0],
        settings.text_model,
        settings.api_host or '(missing)',
        'set' if settings.api_key else 'missing',
    )

    try:
        exit_code = manager.run()

    except KeyboardInterrupt:
        manager.console.interrupted()
        log.warning('session interrupted')
        logging.shutdown()
        os._exit(130)

    log.info(
        'search session end  searched=%s  failed=%s',
        manager.success_count,
        manager.fail_count,
    )

    return exit_code


def unhandled_exception(
    exception_type: type[BaseException],
    exception: BaseException,
    traceback_object: TracebackType | None,
) -> None:
    log.exception(
        'unhandled exception',
        exc_info=(exception_type, exception, traceback_object),
    )

    sys.__excepthook__(exception_type, exception, traceback_object)


if __name__ == '__main__':
    sys.exit(main())
