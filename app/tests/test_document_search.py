from __future__ import annotations

import sys

import pytest

from docx import Document

from pathlib import Path
from typing_extensions import Callable

from data import TranscriptionSettings
from document_search import DocumentSearchManager, document_text_read
from errors import SearchError

from .doubles import SilentConsole


class StubSearcher:
    def __init__(self, findings: str = 'found it', error: Exception | None = None) -> None:
        self.error = error
        self.findings = findings
        self.searched: list[str] = []

    @staticmethod
    def format_search(source_file: Path, query: str, findings: str) -> str:
        return f'{source_file.name}: {findings}'

    def search(self, document_text: str, query: str) -> str:
        self.searched.append(query)

        if self.error:
            raise self.error

        return self.findings

    def search_file_write(self, source_file: Path, query: str, findings: str) -> Path:
        output_file = source_file.with_name(f'{source_file.stem}_search.txt')
        output_file.write_text(findings, encoding='utf-8')

        return output_file


def manager_built(
        settings: TranscriptionSettings,
        console: SilentConsole | None = None,
) -> tuple[DocumentSearchManager, SilentConsole]:
    console = console if console else SilentConsole()

    return DocumentSearchManager(settings, console), console


def test_collect_document_files_globs_the_working_directory(
        monkeypatch: pytest.MonkeyPatch,
        tmp_path: Path,
) -> None:
    (tmp_path / 'report.md').write_text('body', encoding='utf-8')
    (tmp_path / 'notes.txt').write_text('body', encoding='utf-8')
    (tmp_path / 'photo.png').write_bytes(b'image')

    monkeypatch.chdir(tmp_path)
    monkeypatch.setattr(sys, 'argv', ['document_search.py'])

    names = sorted(file.name for file in DocumentSearchManager.collect_document_files())

    assert names == ['notes.txt', 'report.md']


def test_collect_document_files_reads_the_list_file(
        monkeypatch: pytest.MonkeyPatch,
        tmp_path: Path,
) -> None:
    list_file = tmp_path / 'queue.txt'
    lines = f'{tmp_path / "one.txt"}\n\n  {tmp_path / "two.md"}  \n'
    list_file.write_text(lines, encoding='utf-8')

    monkeypatch.setattr(sys, 'argv', ['document_search.py', str(list_file)])

    document_files = DocumentSearchManager.collect_document_files()

    assert [file.name for file in document_files] == ['one.txt', 'two.md']


def test_document_text_read_reads_a_docx(tmp_path: Path) -> None:
    document_file = tmp_path / 'report.docx'
    document = Document()
    _ = document.add_paragraph('first paragraph')
    _ = document.add_paragraph('   ')

    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = 'left'
    table.rows[0].cells[1].text = 'right'
    document.save(str(document_file))

    text = document_text_read(document_file)

    assert 'first paragraph' in text
    assert 'left | right' in text


def test_document_text_read_rejects_an_unsupported_extension(tmp_path: Path) -> None:
    document_file = tmp_path / 'report.pdf'
    document_file.write_bytes(b'pdf')

    with pytest.raises(SearchError) as error:
        _ = document_text_read(document_file)

    assert 'unsupported extension' in str(error.value)


def test_document_text_read_strips_a_byte_order_mark(tmp_path: Path) -> None:
    document_file = tmp_path / 'report.txt'
    document_file.write_bytes(b'\xef\xbb\xbfplain body')

    assert document_text_read(document_file) == 'plain body'


def test_process_counts_a_failure(
        settings_factory: Callable[..., TranscriptionSettings],
        tmp_path: Path,
) -> None:
    document_file = tmp_path / 'report.txt'
    document_file.write_text('body', encoding='utf-8')
    manager, console = manager_built(settings_factory(search_query='renewal date'))

    manager.process(StubSearcher(error=SearchError('model down')), document_file)

    assert manager.fail_count == 1
    assert manager.success_count == 0
    assert console.events_named('failure')[0][2] == 'model down'


def test_process_saves_when_the_console_agrees(
        settings_factory: Callable[..., TranscriptionSettings],
        tmp_path: Path,
) -> None:
    document_file = tmp_path / 'report.txt'
    document_file.write_text('body', encoding='utf-8')
    manager, console = manager_built(settings_factory(search_query='renewal date'), SilentConsole(True))

    manager.process(StubSearcher(), document_file)

    assert manager.success_count == 1
    assert manager.output_files[0].name == 'report_search.txt'
    assert 'search_saved' in console.event_names()


def test_process_skips_saving_when_the_console_declines(
        settings_factory: Callable[..., TranscriptionSettings],
        tmp_path: Path,
) -> None:
    document_file = tmp_path / 'report.txt'
    document_file.write_text('body', encoding='utf-8')
    manager, console = manager_built(settings_factory(search_query='renewal date'), SilentConsole(False))

    manager.process(StubSearcher(), document_file)

    assert manager.output_files == []
    assert 'search_skipped' in console.event_names()


def test_report_result_file_writes_every_output_path(
        settings_factory: Callable[..., TranscriptionSettings],
        tmp_path: Path,
) -> None:
    result_file = tmp_path / 'result.txt'
    settings = settings_factory(result_file=result_file)
    manager, _ = manager_built(settings)
    manager.output_files = [tmp_path / 'one_search.txt', tmp_path / 'two_search.txt']

    manager.report_result_file()

    expected = [str(tmp_path / 'one_search.txt'), str(tmp_path / 'two_search.txt')]

    assert result_file.read_text(encoding='utf-8').splitlines() == expected


def test_report_result_file_writes_nothing_without_outputs(
        settings_factory: Callable[..., TranscriptionSettings],
        tmp_path: Path,
) -> None:
    result_file = tmp_path / 'result.txt'
    manager, _ = manager_built(settings_factory(result_file=result_file))

    manager.report_result_file()

    assert not result_file.exists()


def test_run_reports_a_missing_query(
        monkeypatch: pytest.MonkeyPatch,
        settings_factory: Callable[..., TranscriptionSettings],
) -> None:
    monkeypatch.setattr(sys, 'argv', ['document_search.py'])
    manager, console = manager_built(settings_factory(search_query=''))

    assert manager.run() == 1
    assert console.events_named('error')[0][1] == 'No search request was given.'


def test_run_reports_missing_settings(
        monkeypatch: pytest.MonkeyPatch,
        settings_factory: Callable[..., TranscriptionSettings],
) -> None:
    monkeypatch.setattr(sys, 'argv', ['document_search.py'])
    manager, console = manager_built(settings_factory(api_key=None, search_query='renewal date'))

    assert manager.run() == 1
    assert console.events_named('error')[0][1] == 'No API settings were found.'


def test_run_warns_when_there_is_nothing_to_search(
        monkeypatch: pytest.MonkeyPatch,
        settings_factory: Callable[..., TranscriptionSettings],
) -> None:
    monkeypatch.setattr(sys, 'argv', ['document_search.py'])
    manager, console = manager_built(settings_factory(search_query='renewal date'))

    assert manager.run() == 0
    assert console.events_named('warning')[0][1] == 'No documents to search.'
